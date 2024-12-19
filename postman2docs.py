import json
import os
from docx import Document

# Define file names as variables
input_file = "test.postman_collection.json"
output_folder = "output"
output_file = os.path.join(output_folder, "test.postman_collection.docx")

# Ensure the output folder exists
os.makedirs(output_folder, exist_ok=True)

# Load Postman Collection JSON with UTF-8 encoding
with open(input_file, "r", encoding="utf-8") as file:
    collection = json.load(file)

# Create a new DOCX document
doc = Document()
doc.add_heading('API Documentation', level=1)

# Function to determine data type
def get_datatype(value):
    if isinstance(value, str):
        return "String"
    elif isinstance(value, int):
        return "Integer"
    elif isinstance(value, float):
        return "Float"
    elif isinstance(value, bool):
        return "Boolean"
    elif isinstance(value, dict):
        return "Object"
    elif isinstance(value, list):
        return "Array"
    else:
        return "Unknown"

# Function to create a table for Request Body
def create_request_body_table(body_raw):
    try:
        # Parse JSON string to Python object
        body_data = json.loads(body_raw)
        
        # Add table to the document
        doc.add_heading('Request Body Schema', level=3)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Datatype'

        # Populate table rows
        if isinstance(body_data, dict):  # If the body is a dictionary
            for key, value in body_data.items():
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = f"Description of {key}"  # Placeholder for description
                row_cells[2].text = get_datatype(value)
        elif isinstance(body_data, list):  # If the body is a list
            for idx, item in enumerate(body_data):
                row_cells = table.add_row().cells
                row_cells[0].text = f"Item {idx + 1}"
                row_cells[1].text = f"Description of Item {idx + 1}"  # Placeholder for description
                row_cells[2].text = get_datatype(item)
    except json.JSONDecodeError:
        doc.add_paragraph("Request Body is not a valid JSON.")

# Recursive function to process items
def process_items(items):
    for item in items:
        doc.add_heading(item.get('name', 'Unnamed Endpoint'), level=2)

        # Check if request exists
        if 'request' in item:
            request = item['request']
            
            # Add endpoint and method
            doc.add_paragraph(f"**Endpoint**: {request['url']['raw']}")
            doc.add_paragraph(f"**Method**: {request['method']}")

            # Add headers
            if 'header' in request:
                doc.add_heading('Headers', level=3)
                for header in request['header']:
                    doc.add_paragraph(f"{header['key']}: {header['value']}")

            # Add request body as a table
            if 'body' in request and 'raw' in request['body']:
                create_request_body_table(request['body']['raw'])

            # Add request body
            if 'body' in request and 'raw' in request['body']:
                doc.add_heading('Request Body', level=3)
                doc.add_paragraph(request['body']['raw'])

        # Handle sub-items (if nested)
        if 'item' in item:
            process_items(item['item'])

# Process top-level items
process_items(collection.get('item', []))

# Save the DOCX document
doc.save(output_file)

print(f"Documentation generated and saved in {output_file}")
